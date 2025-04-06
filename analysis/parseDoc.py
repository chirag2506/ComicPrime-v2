from docx import Document
from copy import deepcopy
from datetime import datetime
from utilities import writeJson

DOCUMENT_NAME = 'files/docx/Comics  (1939-Feb 2021).docx'


document = Document(DOCUMENT_NAME)

startGetting = False
# index = 35147 # total num comics (1939-Feb 2021)
index = 0
comicsPerMonth = {}

while(index < len(document.paragraphs) - 1):
    # if(index>100):
    #     break
    para = document.paragraphs[index]
    text = para.text
    validMonth = True
    try:
        validCheck = datetime.strptime(text, '%B, %Y')
        startGetting = True
    except Exception as e:
        validMonth = False
    if(startGetting == True):
        if(validMonth == True):
            # found a new month
            print("*"*100)
            month = deepcopy(text)
            comicsPerMonth[month] = {}
            numberInMonthText = document.paragraphs[index + 1]
            numComics = int(numberInMonthText.text[38:]) # "Number of comics published in this month: xx" line
            firstComic = index + 2 # skipping "Number of comics published in this month: xx" line
            finalComic = index + numComics + 1 # skipping line and going to final comic in the month
            print("Comics published in: {} ({})".format(month, numComics))
            comicsPerMonth[month]["totalComics"] = numComics
            comicsInMonth = []
            for index in range(firstComic, finalComic+1):
                para = document.paragraphs[index]
                text = para.text
                comicsInMonth.append(text)
            comicsPerMonth[month]["comics"] = comicsInMonth
    index += 1

writeJson("files/json/comics.json", comicsPerMonth)
