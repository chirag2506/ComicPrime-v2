from bs4 import BeautifulSoup as soup
from urllib.request import urlopen as uReq
from datetime import date
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import calendar

totalComics = 0

#Document

document = Document()
document.add_heading('List of Marvel Comics Published Through The Years', 0)
document.add_picture('files/images/logo.jpg', width=Inches(4))
lastParagraph = document.paragraphs[-1]
lastParagraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#Scraping

def monthYearIter( startMonth, startYear, endMonth, endYear ):
    listOfMonths = []
    ymStart= 12*startYear + startMonth - 1
    ymEnd= 12*endYear + endMonth - 1
    for ym in range( ymStart, ymEnd ):
        y, m = divmod( ym, 12 )
        listOfMonths.append([y, m+1])
    return listOfMonths


currentYear = date.today().year
currentMonth = date.today().month
# listOfMonths = monthYearIter(11,1961,currentMonth,currentYear)
listOfMonths = monthYearIter(12,1939,currentMonth,currentYear)


for months in range(0, len(listOfMonths)):

    thisYear = listOfMonths[months][0]
    thisMonth = listOfMonths[months][1]
    thisMonth = "{0:0=2d}".format(thisMonth) #To convert to 2-digit (1 --> 01)

    #DOCUMENT
    document.add_heading(calendar.month_name[int(thisMonth)] + ', ' + str(thisYear), level=3)

    #SCRAPE
    url = 'https://www.marvel.com/comics/calendar/month/' + str(thisYear) + '-' + str(thisMonth) + '-01'
    print(url)
    myUrl = url

    uClient = uReq(myUrl)
    pageHtml = uClient.read()
    uClient.close()
    pageSoup = soup(pageHtml, "html.parser")

    containers = pageSoup.findAll("div", {"class": "row-item comic-item"})
    thisMonthNumComics = len(containers)
    #print(len(containers))
    totalComics = totalComics + len(containers)

    #DOCUMENT
    p = document.add_paragraph('Number of comics published this month: ')
    p.add_run(str(thisMonthNumComics)).bold = True


    for comic in range(0, len(containers)):
        container = containers[comic]
        comicName = container.div.next_sibling.next_sibling.h5.a.text
        #print(comicName.strip())
        document.add_paragraph(
            comicName.strip(), style='List Bullet'
        )

    document.save('files/docx/Comics.docx')

comicName = document.paragraphs[2]
comicName.insert_paragraph_before('MARVEL HAS PUBLISHED ABOUT ' + str(totalComics) + ' COMICS TILL DATE')
document.save('files/docx/Comics.docx')
