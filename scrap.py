from bs4 import BeautifulSoup as soup
from urllib.request import urlopen as uReq
from datetime import date
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import calendar

total_comics = 0

#Document

document = Document()
document.add_heading('List of Marvel Comics Published Through The Years', 0)
document.add_picture('logo.jpg', width=Inches(4))
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#Scraping

def month_year_iter( start_month, start_year, end_month, end_year ):
    list_of_months = []
    ym_start= 12*start_year + start_month - 1
    ym_end= 12*end_year + end_month - 1
    for ym in range( ym_start, ym_end ):
        y, m = divmod( ym, 12 )
        list_of_months.append([y, m+1])
    return list_of_months


current_year = date.today().year
current_month = date.today().month
# list_of_months = month_year_iter(11,1961,current_month,current_year)
list_of_months = month_year_iter(12,1939,current_month,current_year)


for months in range(0, len(list_of_months)):

    this_year = list_of_months[months][0]
    this_month = list_of_months[months][1]
    this_month = "{0:0=2d}".format(this_month) #To convert to 2-digit (1 --> 01)

    #DOCUMENT
    document.add_heading(calendar.month_name[int(this_month)] + ', ' + str(this_year), level=3)

    #SCRAPE
    url = 'https://www.marvel.com/comics/calendar/month/' + str(this_year) + '-' + str(this_month) + '-01'
    print(url)
    my_url = url

    uClient = uReq(my_url)
    page_html = uClient.read()
    uClient.close()
    page_soup = soup(page_html, "html.parser")

    containers = page_soup.findAll("div", {"class": "row-item comic-item"})
    this_month_num_comics = len(containers)
    #print(len(containers))
    total_comics = total_comics + len(containers)

    #DOCUMENT
    p = document.add_paragraph('Number of comics published this month: ')
    p.add_run(str(this_month_num_comics)).bold = True


    for comic in range(0, len(containers)):
        container = containers[comic]
        comic_name = container.div.next_sibling.next_sibling.h5.a.text
        #print(comic_name.strip())
        document.add_paragraph(
            comic_name.strip(), style='List Bullet'
        )

    document.save('Comics.docx')

first_month = document.paragraphs[2]
first_month.insert_paragraph_before('MARVEL HAS PUBLISHED ABOUT ' + str(total_comics) + ' COMICS TILL DATE')
document.save('Comics.docx')
