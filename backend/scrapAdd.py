from bs4 import BeautifulSoup as soup
from urllib.request import urlopen as uReq
from docx import Document

DOCUMENT_NAME = 'ComicsNew.docx'
# DOCUMENT_NAME = 'Comics to read (1939-Feb 2021).docx'
# DOCUMENT_NAME = 'Comics  (1939-Feb 2021).docx'
URL = 'https://marvel.fandom.com/wiki/Western_Gunfighters_Vol_2'
URL = 'https://marvel.fandom.com/wiki/Marvel_Spotlight_Vol_1'

document = Document(DOCUMENT_NAME)

paras = []
for paragraph in document.paragraphs:
        paras.append(paragraph.text)

myUrl = URL
uClient = uReq(myUrl)
pageHtml = uClient.read()
uClient.close()
pageSoup = soup(pageHtml, "html.parser")

containers = pageSoup.findAll("div", {"class": "wikia-gallery-item"})
comicsInThisVolume = len(containers)
print("COMICS IN THIS VOLUME: " + str(comicsInThisVolume))

for comic in range(0, len(containers)):
    container = containers[comic]
    comicName = container.div.next_sibling.div.span.a.text
    releaseAndCover = container.div.next_sibling.div.span.next_sibling.text
    coverDate = releaseAndCover[releaseAndCover.find("Cover date: ")+12 : ]
    indexOfCover = paras.index(coverDate)
    #changing number of comics
    numberInMonth = document.paragraphs[indexOfCover + 1 + comic] # "+ comic" done because new line is being added after every iteration
    numberInMonthInNewText = int(numberInMonth.text[38:]) + 1
    numberInMonth.text = "Number of comics published this month: "
    numberInMonth.add_run(str(numberInMonthInNewText)).bold=True
    #adding the comic
    document.paragraphs[indexOfCover+2+comic].insert_paragraph_before(
        comicName.strip(), style='List Bullet'
    )

    document.save(DOCUMENT_NAME)

document.save(DOCUMENT_NAME)